# backend/app.py
from fastapi import FastAPI, Request, HTTPException
from fastapi.responses import JSONResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from .services.vector_store import ChromaVectorStore
from backend.agents.internal_vector_agent import InternalVectorAgent
from backend.agents import planner, external_web_agent, create_report_agent, plagiarism_agent
from backend.services import session_service
import os
import tempfile
import json
import textwrap
from fpdf import FPDF
import unicodedata
import re
import logging
from backend.agents.create_report_agent import Source # Import the Source dataclass

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[logging.FileHandler("backend_log.txt", encoding="utf-8"), logging.StreamHandler()]
)

# --- APPLICATION BOOT SEQUENCE ---
print("[BOOT] Initializing application...")
VECTOR_STORE = ChromaVectorStore()
collection_count = VECTOR_STORE.collection.count()
print(f"[BOOT] Connected to Chroma store. Found {collection_count} existing document chunks.")
if collection_count == 0:
    print("[BOOT] WARNING: The vector store is empty. Run 'python ingest.py' to add documents.")
INTERNAL_AGENT = InternalVectorAgent(VECTOR_STORE)
print("[BOOT] Application is ready to serve requests.")
# --- Boot Complete ---

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], allow_credentials=True, allow_methods=["*"], allow_headers=["*"]
)

def to_ascii_safe(text: str) -> str:
    text = (text.replace("'", "'").replace("“", '"').replace("”", '"').replace("—", "-").replace("–", "-"))
    text = unicodedata.normalize("NFKD", text).encode("ascii", "ignore").decode("ascii")
    text = re.sub(r"[ \t]+", " ", text)
    return text

@app.post("/refresh_internal_index")
async def refresh_internal_index():
    logging.error("[REFRESH] This endpoint is disabled. Run 'python ingest.py' from the terminal to refresh the index.")
    raise HTTPException(status_code=403, detail="Manual refresh is disabled. Use the 'ingest.py' script.")

@app.post("/plan/")
async def generate_plan(request: Request):
    data = await request.json()
    topic = (data.get("topic") or "").strip()
    if not topic:
        raise HTTPException(status_code=400, detail="Missing 'topic'")
    plan = planner.generate_research_plan(topic, INTERNAL_AGENT)
    sess = session_service.session_manager.get_session(request)
    sess['plan'] = plan
    sess['topic'] = topic
    return JSONResponse({"plan": plan})


# In app.py - Add these imports at the top
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import tempfile
import os

@app.post("/execute/")
async def execute_plan(request: Request):
    sess = session_service.session_manager.get_session(request)
    plan_from_session = sess.get('plan')
    data = await request.json()
    plan_from_body = data.get('plan')
    plan = plan_from_body or plan_from_session

    if not plan:
        raise HTTPException(status_code=404, detail="No plan found. Please generate a plan first.")

    if isinstance(plan, str):
        try: plan = json.loads(plan)
        except Exception: raise HTTPException(status_code=400, detail="Plan must be valid JSON")

    if not isinstance(plan, dict) or 'steps' not in plan or 'topic' not in plan:
        raise HTTPException(status_code=422, detail="Plan JSON missing 'topic' or 'steps'")

    topic, steps = plan['topic'], plan['steps'][:6]
    raw_sources = []

    for step in steps:
        query = (step.get('query') or "").strip()
        if not query: continue

        try:
            # Retrieve internal documents and create Source objects
            internal_results = INTERNAL_AGENT.retrieve(query)
            for r in internal_results:
                meta = r.get("metadata", {})
                raw_sources.append(
                    Source(
                        content=r["text"],
                        citation=f"({meta.get('file_name', 'Internal Document')} p.{meta.get('page', 'N/A')})".strip(),
                        source_type="internal",
                        title=meta.get('file_name', 'Internal Document')
                    )
                )

            # Retrieve external documents and create Source objects  
            external_results = external_web_agent.web_search(query)
            for r in external_results:
                if r.get('snippet'):
                    url = r.get('url', '') or r.get('link', '')
                    raw_sources.append(
                        Source(
                            content=r.get('snippet', ''),
                            citation=f"({r.get('title', 'External Source')})",
                            source_type="external",
                            title=r.get('title', 'External Source'),
                            url=url
                        )
                    )
        except Exception as e:
            logging.error(f"Error processing query '{query}': {e}")
            continue

    # Ensure source diversity and quality
    sources = create_report_agent.ensure_source_diversity(raw_sources)
    logging.info(f"Using {len(sources)} diverse sources for report generation")

    # Generate the full report text from the collected sources
    full_text = create_report_agent.create_research_report(sources, topic)

    # Validate content quality
    validation = create_report_agent.validate_content_before_pdf(full_text, sources)
    logging.info(f"Content validation score: {validation['quality_score']}")
    
    if validation["quality_score"] < 0.6:
        logging.warning(f"Low quality score: {validation['quality_score']}. Regenerating with enhanced settings.")
        full_text = create_report_agent.create_research_report(
            sources, topic, enhanced_rephrasing=True
        )

    # Create professional Word document
    doc_path = create_professional_word_doc(full_text, topic)
    
    # Create a safe filename
    safe_filename = f"research_report_{topic.replace(' ', '_').lower()}.docx"
    safe_filename = re.sub(r'[^\w\-_\.]', '', safe_filename)
    
    return FileResponse(
        path=doc_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=safe_filename,
        headers={"Content-Disposition": f"attachment; filename={safe_filename}"}
    )


def create_professional_word_doc(full_text: str, topic: str) -> str:
    """Create a professionally formatted Word document"""
    
    doc = Document()
    
    # Set up document styles
    setup_document_styles(doc)
    
    # Add title page
    add_title_page(doc, topic)
    
    # Add page break before content
    doc.add_page_break()
    
    # Process content
    add_formatted_content(doc, full_text)
    
    # Save to temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    temp_file.close()
    
    return temp_file.name

def setup_document_styles(doc):
    """Set up custom styles for the document"""
    
    styles = doc.styles
    
    # Create title style
    try:
        title_style = styles['Title']
    except KeyError:
        title_style = styles.add_style('Title', WD_STYLE_TYPE.PARAGRAPH)
    
    title_style.font.name = 'Arial'
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(12)
    
    # Create heading style
    try:
        heading_style = styles['Heading 1']
    except KeyError:
        heading_style = styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
    
    heading_style.font.name = 'Arial'
    heading_style.font.size = Pt(16)
    heading_style.font.bold = True
    heading_style.paragraph_format.space_before = Pt(12)
    heading_style.paragraph_format.space_after = Pt(6)
    
    # Create normal style
    try:
        normal_style = styles['Normal']
    except KeyError:
        normal_style = styles.add_style('Normal', WD_STYLE_TYPE.PARAGRAPH)
    
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(11)
    normal_style.paragraph_format.space_after = Pt(6)
    normal_style.paragraph_format.line_spacing = 1.15

def add_title_page(doc, topic):
    """Add a professional title page"""
    
    # Add title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(topic.upper())
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(24)
    title_run.bold = True
    
    # Add some space
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Add subtitle
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run('Research Report')
    subtitle_run.font.name = 'Arial'
    subtitle_run.font.size = Pt(16)
    
    # Add more space
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Add date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run('Generated on: August 11, 2025')
    date_run.font.name = 'Arial'
    date_run.font.size = Pt(12)

def add_formatted_content(doc, full_text):
    """Add formatted content to the document"""
    
    lines = full_text.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
        
        # Check if it's a section header (ALL CAPS)
        if line.isupper() and len(line) > 3 and not line.startswith('['):
            # Add section header
            header_para = doc.add_paragraph()
            header_run = header_para.add_run(line)
            header_run.font.name = 'Arial'
            header_run.font.size = Pt(16)
            header_run.bold = True
            header_para.paragraph_format.space_before = Pt(12)
            header_para.paragraph_format.space_after = Pt(6)
            
            # Skip the dashes line if it follows
            if i + 1 < len(lines) and lines[i + 1].strip().startswith('---'):
                i += 1
        
        elif line.startswith('---'):
            # Skip dash lines
            pass
        
        else:
            # Regular paragraph with citation handling
            add_paragraph_with_citations(doc, line)
        
        i += 1

def add_paragraph_with_citations(doc, text):
    """Add a paragraph with proper citation formatting"""
    
    if not text.strip():
        return
    
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.15
    
    # Split text by citations
    parts = re.split(r'(\[\d+\])', text)
    
    for part in parts:
        if not part:
            continue
        
        if re.match(r'^\[\d+\]$', part):
            # Add citation as superscript
            citation_run = para.add_run(part)
            citation_run.font.name = 'Arial'
            citation_run.font.size = Pt(9)
            citation_run.font.superscript = True
        else:
            # Add regular text
            text_run = para.add_run(part)
            text_run.font.name = 'Arial'
            text_run.font.size = Pt(11)
